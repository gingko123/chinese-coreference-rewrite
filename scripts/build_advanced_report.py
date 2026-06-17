from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
BASE_REPORT = ROOT / "reports" / "期末报告_中文指代消解与句子改写系统.docx"
OUTPUT_REPORT = ROOT / "reports" / "期末报告_中文指代消解与句子改写系统_进阶版.docx"


ABLATION_ROWS = [
    ("test", "完整规则", "90/90", "1.0000", "1.0000", "1.0000", "+0.0000"),
    ("test", "去掉距离规则", "84/90", "0.9333", "0.9333", "0.9333", "-0.0667"),
    ("test", "去掉位置规则", "74/90", "0.8222", "0.8222", "0.8222", "-0.1778"),
    ("test", "去掉类型规则", "44/90", "0.4889", "0.4889", "0.4889", "-0.5111"),
    ("test", "去掉性别规则", "90/90", "1.0000", "1.0000", "1.0000", "+0.0000"),
    ("real_corpus", "完整规则", "44/46", "0.9565", "0.9565", "0.9565", "+0.0000"),
    ("real_corpus", "去掉距离规则", "32/46", "0.6957", "0.6957", "0.6957", "-0.2608"),
    ("real_corpus", "去掉位置规则", "43/46", "0.9348", "0.9348", "0.9348", "-0.0217"),
    ("real_corpus", "去掉类型规则", "39/46", "0.8478", "0.8478", "0.8478", "-0.1087"),
    ("real_corpus", "去掉性别规则", "43/46", "0.9348", "0.9348", "0.9348", "-0.0217"),
]


def set_cell_shading(cell, fill):
    """Set table cell background color."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_text(cell, text, bold=False):
    """Write text into a table cell with consistent font."""
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9)


def add_body_paragraph(doc, text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)
    return paragraph


def add_code_paragraph(doc, text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    return paragraph


def add_heading(doc, text, level):
    paragraph = doc.add_heading(level=level)
    run = paragraph.add_run(text)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    return paragraph


def add_ablation_table(doc):
    headers = ["数据集", "实验组", "正确/真实", "Precision", "Recall", "F1", "ΔF1"]
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        # Some Chinese Word templates do not include this built-in English style name.
        pass
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True)
        set_cell_shading(table.rows[0].cells[index], "D9EAF7")

    for row_data in ABLATION_ROWS:
        cells = table.add_row().cells
        for index, value in enumerate(row_data):
            set_cell_text(cells[index], value)


def build_report():
    if not BASE_REPORT.exists():
        raise FileNotFoundError(f"Base report not found: {BASE_REPORT}")

    doc = Document(BASE_REPORT)
    doc.paragraphs[-1].add_run().add_break(WD_BREAK.PAGE)

    add_heading(doc, "九、进阶实验与应用扩展", 1)
    add_body_paragraph(
        doc,
        "为增强项目的实验完整性，本项目在原有“规则指代消解 + 文本改写 + Streamlit 可视化”的基础上，"
        "增加了消融实验、错误分析、BERT 语义匹配增强方案以及面向问答场景的应用设计。"
        "这些内容能够说明系统不仅可以运行，还能够被评估、被分析，并具备进一步扩展为真实 NLP 应用的空间。",
    )

    add_heading(doc, "9.1 消融实验目的", 2)
    add_body_paragraph(
        doc,
        "当前系统使用距离、位置、实体类型和性别四类规则信号对候选先行词进行打分。"
        "消融实验的目的，是在保持数据集和评价方式不变的条件下，依次去掉某一类规则，"
        "观察 F1 值下降幅度，从而判断不同规则对指代消解效果的贡献。"
    )
    add_code_paragraph(doc, "final_score = distance_score + position_score + type_score + gender_score")

    add_heading(doc, "9.2 实验设置", 2)
    add_body_paragraph(
        doc,
        "实验使用两个数据集：test.json 代表构造测试集，样本模式较规范；real_corpus.json 代表真实语料风格数据，"
        "文本更自然，也更容易暴露规则方法的局限。评价指标采用 Precision、Recall 和 F1，"
        "同时记录与完整规则相比的 F1 变化。"
    )

    add_heading(doc, "9.3 消融实验结果", 2)
    add_ablation_table(doc)

    add_heading(doc, "9.4 结果分析", 2)
    add_body_paragraph(
        doc,
        "从构造测试集看，完整规则达到 1.0000 的 F1。去掉类型规则后 F1 降至 0.4889，"
        "说明人物、物体、组织、地点等实体类型匹配是当前规则系统中最关键的判断依据。"
        "去掉距离和位置规则也会造成下降，但影响小于类型规则。"
    )
    add_body_paragraph(
        doc,
        "从真实语料风格数据看，完整规则 F1 为 0.9565。去掉距离规则后 F1 降至 0.6957，"
        "说明在更自然的文本中，候选先行词与代词之间的局部距离仍然是非常重要的线索。"
        "同时，真实语料中更容易出现多个候选实体竞争、地点名与物体名混淆、抽象事件指代等问题。"
    )
    add_body_paragraph(
        doc,
        "性别规则在当前数据上的影响较小，主要原因是数据集中“他/她”歧义样本数量有限。"
        "后续可以继续补充包含多个人物、性别冲突和跨句指代的样本，使评估更加全面。"
    )

    add_heading(doc, "9.5 BERT 语义匹配增强方案", 2)
    add_body_paragraph(
        doc,
        "规则方法可解释性强，但对复杂语义理解能力有限。后续可以引入中文 BERT、MacBERT 或 sentence-transformers，"
        "对代词上下文和候选先行词上下文分别编码，计算语义相似度 semantic_score，再与原有规则分数融合。"
    )
    add_code_paragraph(doc, "final_score = 0.7 * rule_score + 0.3 * semantic_score")
    add_body_paragraph(
        doc,
        "例如在“上海市启动了城市更新计划，该城市希望它改善居民生活环境”这类句子中，"
        "单纯依靠距离和类型规则容易把“它”错误地指向较近的词；加入语义匹配后，"
        "系统可以利用上下文判断“改善居民生活环境”更可能对应“城市更新计划”这一事件或方案。"
    )

    add_heading(doc, "9.6 面向问答场景的应用", 2)
    add_body_paragraph(
        doc,
        "指代消解还可以作为问答系统的预处理模块。用户问题中经常包含“他、她、它、该公司、该城市”等表达，"
        "如果不先完成指代消解，后续答案抽取模块可能无法准确定位答案。"
    )
    add_code_paragraph(doc, "输入：小明把书给了小红，因为她需要它。谁需要书？")
    add_code_paragraph(doc, "改写：小明把书给了小红，因为小红需要书。")
    add_code_paragraph(doc, "答案：小红")
    add_body_paragraph(
        doc,
        "因此，本项目不仅是一个文本改写演示系统，也可以扩展为问答、摘要、信息抽取等下游任务的基础模块。"
        "这使项目具有更明确的应用价值，也能让课程报告从“功能展示”提升到“实验分析 + 应用设计”的层次。"
    )

    doc.save(OUTPUT_REPORT)
    return OUTPUT_REPORT


if __name__ == "__main__":
    print(build_report())
