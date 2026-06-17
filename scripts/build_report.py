from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = Path(
    r"C:\Users\30657\Documents\xwechat_files\wxid_843vwkgrw2gj22_03a2\msg\file\2026-06\期末报告模板.docx"
)
OUT_DIR = ROOT / "reports"
OUT_PATH = OUT_DIR / "期末报告_中文指代消解与句子改写系统.docx"


def set_run_font(run, size=12, bold=False, color=None, font="宋体"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_paragraph(paragraph, size=12, bold=False, align=None, line_spacing=1.5):
    paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.paragraph_format.space_after = Pt(6)
    if align is not None:
        paragraph.alignment = align
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold)


def clear_body(doc):
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def add_para(doc, text="", size=12, bold=False, align=None, style=None, before=0, after=6):
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return paragraph


def add_heading(doc, text, level=1):
    sizes = {1: 16, 2: 15, 3: 14}
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run, size=sizes[level], bold=True)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style=None)
    paragraph.paragraph_format.left_indent = Cm(0.74)
    paragraph.paragraph_format.first_line_indent = Cm(-0.32)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run("• " + text)
    set_run_font(run, size=12)
    return paragraph


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.line_spacing = 1.2
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "666666")


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        set_cell_text(header_cells[index], header, bold=True)
        shade_cell(header_cells[index], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], str(value))
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Cm(width)
    doc.add_paragraph()
    return table


def add_code_block(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.6)
    paragraph.paragraph_format.line_spacing = 1.1
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, font="Consolas")
    return paragraph


def build_report():
    OUT_DIR.mkdir(exist_ok=True)
    doc = Document(str(TEMPLATE))
    clear_body(doc)

    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)

    add_para(doc, "华东师范大学", size=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=24)
    add_para(doc, "《自然语言处理》", size=20, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    add_para(doc, "期末报告", size=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=30)
    add_para(
        doc,
        "标题：中文指代消解与句子改写系统",
        size=15,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=18,
    )
    add_para(doc, "小组成员：", size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_para(doc, "姓名：________   学号：________   学院：________", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "姓名：________   学号：________   学院：________", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "2026年6月12日", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, before=36)

    doc.add_page_break()

    add_heading(doc, "一、项目背景", 1)
    add_heading(doc, "1.1 问题描述与实际需求", 2)
    add_para(
        doc,
        "在自然语言文本中，代词和指代表达大量存在。例如“她”“它”“该公司”“这个项目”等表达本身并不直接给出完整实体，"
        "读者需要结合上下文判断其真实指向。对于机器而言，如果不能正确理解这些指代关系，就会影响文本理解、问答系统、信息抽取、"
        "文本摘要和智能写作等下游任务的效果。",
    )
    add_para(
        doc,
        "本项目围绕中文指代消解与句子改写展开，目标是识别文本中的代词或指代表达，判断其对应的先行实体，并将原句改写为指代更明确的表达。"
        "例如输入“小明把书给了小红，因为她需要它。”，系统应输出“她 -> 小红”“它 -> 书”，并生成“小明把书给了小红，因为小红需要书。”",
    )
    add_heading(doc, "1.2 应用价值", 2)
    add_bullet(doc, "提升文本理解能力：帮助系统显式表示句子中隐含的实体关系。")
    add_bullet(doc, "辅助阅读与写作：将含代词的句子改写为更明确的表达，便于阅读和检查。")
    add_bullet(doc, "服务下游 NLP 任务：为问答、摘要、信息抽取等任务提供更清晰的实体链信息。")

    add_heading(doc, "二、主要功能", 1)
    add_heading(doc, "2.1 系统核心功能", 2)
    add_bullet(doc, "中文文本输入：用户可以输入一段中文文本，也可以选择系统内置示例。")
    add_bullet(doc, "实体与代词识别：系统识别人名、组织、物体名词以及人称代词、指示代词。")
    add_bullet(doc, "候选先行词生成：针对每个代词，筛选其前文或短语内部的候选实体。")
    add_bullet(doc, "指代关系判断：根据距离、类型、性别、位置等规则为候选实体打分。")
    add_bullet(doc, "文本改写：将已消解代词替换为对应实体，生成更明确的句子。")
    add_bullet(doc, "数据集评估：支持 demo、dev、test 数据集的 Accuracy、Precision、Recall 和 F1 评估。")
    add_bullet(doc, "错误分析：展示漏判、误判和预测关系，方便分析系统局限。")

    add_heading(doc, "2.2 功能模块划分", 2)
    add_table(
        doc,
        ["模块", "对应文件", "主要职责"],
        [
            ["文本预处理", "src/preprocess.py", "分句与简单 token 化，为后续分析提供基础结构。"],
            ["Mention 抽取", "src/mention_extractor.py", "抽取候选实体、代词和指示短语。"],
            ["指代消解", "src/resolver.py", "对候选先行词进行规则打分并输出最佳指代关系。"],
            ["文本改写", "src/rewriter.py", "根据消解结果替换代词，生成改写文本。"],
            ["实验评估", "src/evaluator.py", "计算 Accuracy、Precision、Recall、F1 并输出错误样本。"],
            ["可视化界面", "app.py", "使用 Streamlit 展示单文本分析和数据集评估结果。"],
        ],
        widths=[3.0, 4.2, 8.0],
    )
    add_para(doc, "表 1 系统功能模块划分")

    add_heading(doc, "三、技术方案", 1)
    add_heading(doc, "3.1 技术选型", 2)
    add_table(
        doc,
        ["技术", "用途", "选择理由"],
        [
            ["Python", "核心开发语言", "生态成熟，适合 NLP 原型系统开发。"],
            ["Streamlit", "Web 可视化 Demo", "开发成本低，可以快速展示输入、输出、表格和指标。"],
            ["规则方法", "指代关系判断", "可解释性强，适合课程项目和小规模数据集。"],
            ["JSON 数据集", "样本存储", "结构清晰，便于人工标注、读取和扩展。"],
            ["python-docx", "报告生成", "便于按照 Word 模板生成课程报告。"],
        ],
        widths=[3.2, 4.0, 8.0],
    )
    add_para(doc, "表 2 技术选型与理由")

    add_heading(doc, "3.2 系统架构设计", 2)
    add_para(doc, "系统整体采用“输入文本 -> 语言单元识别 -> 候选生成 -> 规则打分 -> 改写与展示”的流水线结构。")
    add_code_block(
        doc,
        "用户输入文本\n"
        "    ↓\n"
        "实体与代词抽取\n"
        "    ↓\n"
        "候选先行词生成\n"
        "    ↓\n"
        "规则打分与排序\n"
        "    ↓\n"
        "指代关系输出\n"
        "    ↓\n"
        "文本改写与 Streamlit 可视化",
    )
    add_para(doc, "图 1 系统处理流程图")

    add_heading(doc, "3.3 数据集设计", 2)
    add_para(
        doc,
        "本项目构建了一个小规模中文指代消解数据集，共 50 条样本、86 个指代关系。样本覆盖人物指代、物体指代、组织指代、"
        "指示短语和多候选实体等类型。数据采用 JSON 格式，每条样本包含原文、指代关系和改写文本。",
    )
    add_table(
        doc,
        ["数据集", "样本数", "用途"],
        [
            ["demo.json", "10", "用于页面示例和课堂演示。"],
            ["dev.json", "20", "用于规则调试和开发阶段验证。"],
            ["test.json", "20", "用于最终实验评估。"],
            ["samples.json", "50", "完整样本备份。"],
        ],
        widths=[3.5, 2.5, 9.0],
    )
    add_para(doc, "表 3 数据集划分")

    add_heading(doc, "四、关键算法与模型原理说明", 1)
    add_heading(doc, "4.1 Mention 抽取", 2)
    add_para(
        doc,
        "系统首先从文本中抽取两类 mention：一类是候选先行实体，包括人名、组织名和物体名词；另一类是需要消解的代词或指代表达。"
        "当前版本采用词表和规则进行识别，例如将“小明”“李娜”等识别为人物实体，将“书”“产品”“系统”等识别为物体实体，"
        "将“他”“她”“它”“该公司”“该系统”“这些作业”等识别为代词或指示短语。",
    )
    add_heading(doc, "4.2 候选先行词生成", 2)
    add_para(
        doc,
        "对于每个待消解代词，系统优先选择出现在代词之前的实体作为候选先行词。对于“这个项目”“该系统”等指示短语，"
        "系统也允许短语内部的核心名词作为候选，以处理“公司开发了项目，这个项目已经进入测试阶段”这类样本。",
    )
    add_heading(doc, "4.3 规则打分机制", 2)
    add_para(doc, "候选实体的最终分数由多个规则信号累加得到，核心因素如下：")
    add_bullet(doc, "距离得分：候选实体距离代词越近，得分越高。")
    add_bullet(doc, "位置得分：候选实体通常应出现在代词之前。")
    add_bullet(doc, "类型得分：人物代词优先匹配人物实体，物体代词优先匹配物体实体，组织指代表达优先匹配组织实体。")
    add_bullet(doc, "性别得分：当人物实体存在性别线索时，与“他”“她”匹配的候选得分更高。")
    add_code_block(
        doc,
        "score = distance_score + position_score + type_score + gender_score",
    )
    add_para(doc, "公式 1 候选先行词规则打分公式")

    add_heading(doc, "4.4 实验结果", 2)
    add_para(
        doc,
        "系统在当前自建数据集上进行关系级别评估。若预测出的“代词 -> 先行词”关系与人工标注一致，则记为正确。"
        "实验结果如下：",
    )
    add_table(
        doc,
        ["数据集", "真实关系数", "预测关系数", "正确数", "Accuracy", "Precision", "Recall", "F1"],
        [
            ["demo", "20", "20", "20", "100%", "100%", "100%", "100%"],
            ["dev", "35", "35", "35", "100%", "100%", "100%", "100%"],
            ["test", "31", "31", "31", "100%", "100%", "100%", "100%"],
            ["samples", "86", "86", "86", "100%", "100%", "100%", "100%"],
        ],
        widths=[2.0, 2.0, 2.0, 1.8, 1.8, 2.0, 1.8, 1.8],
    )
    add_para(doc, "表 4 当前规则系统实验结果")
    add_para(
        doc,
        "需要说明的是，当前测试集规模较小，且数据样本主要围绕规则系统可处理的典型场景构建，因此实验结果更适合作为原型系统的功能验证，"
        "不能代表系统在开放真实语料上的泛化性能。",
    )

    add_heading(doc, "五、界面展示", 1)
    add_heading(doc, "5.1 单文本分析页面", 2)
    add_para(
        doc,
        "系统首页提供“单文本分析”标签页。用户可以输入中文文本，点击“开始分析”后，页面会展示候选实体、代词、指代关系、候选实体得分以及改写结果。"
        "其中实体和代词在原文中使用不同颜色高亮，便于观察系统的中间结果。",
    )
    add_para(doc, "图 2 单文本分析页面示意（可在最终版中替换为实际截图）")
    add_heading(doc, "5.2 数据集评估页面", 2)
    add_para(
        doc,
        "“数据集评估”标签页支持选择 demo、dev、test 或完整样本集。页面以指标卡形式展示 Accuracy、Precision、Recall、F1，"
        "并在存在错误时展示错误样本表，包括原文、真实关系、预测关系、漏判和误判。",
    )
    add_para(doc, "图 3 数据集评估页面示意（可在最终版中替换为实际截图）")

    add_heading(doc, "六、思考与总结", 1)
    add_heading(doc, "6.1 开发过程中遇到的问题与解决方案", 2)
    add_bullet(doc, "中文姓名和组织名边界容易识别错误。解决方案是收紧姓名规则，并加入常见姓名词表。")
    add_bullet(doc, "“该系统”“该报告”“这些作业”等指示短语类型较多。解决方案是扩展代词词表，并统一映射到人物、物体或组织类型。")
    add_bullet(doc, "Streamlit 在普通 Python 模式下会出现 missing ScriptRunContext 警告。解决方案是使用 streamlit run app.py 启动。")
    add_bullet(doc, "GitHub 远程仓库预置 README 导致首次 push 冲突。解决方案是 fetch 后使用 force-with-lease 推送本地完整项目。")

    add_heading(doc, "6.2 项目局限性", 2)
    add_bullet(doc, "当前实体识别主要依赖规则和词表，面对开放领域文本时召回率有限。")
    add_bullet(doc, "规则方法缺少深层语义理解，难以处理需要常识推理或复杂句法分析的指代。")
    add_bullet(doc, "改写模块采用直接替换策略，生成结果有时会略显重复，不如生成式模型自然。")
    add_bullet(doc, "数据集规模较小，评估结果主要反映原型功能，不足以证明真实场景泛化能力。")

    add_heading(doc, "6.3 未来改进方向", 2)
    add_bullet(doc, "接入 HanLP 或 LTP，提高中文分词、词性标注和命名实体识别能力。")
    add_bullet(doc, "引入 BERT 或 sentence-transformers，计算候选实体与代词上下文的语义相似度。")
    add_bullet(doc, "扩展真实新闻、故事和问答文本样本，构建更有挑战性的测试集。")
    add_bullet(doc, "优化改写策略，减少重复表达，提升改写文本的自然度。")

    add_heading(doc, "七、小组分工与贡献", 1)
    add_table(
        doc,
        ["成员", "主要工作", "贡献说明"],
        [
            ["成员 1", "项目设计、核心算法、报告撰写", "负责整体方案、规则打分机制和实验分析。"],
            ["成员 2", "数据标注、界面实现、测试验证", "负责样本整理、Streamlit 页面和演示材料。"],
        ],
        widths=[3.0, 5.0, 7.0],
    )
    add_para(doc, "表 5 小组分工与贡献")

    add_heading(doc, "八、参考资料", 1)
    add_para(doc, "[1] Hugging Face. neuralcoref: Coreference Resolution in spaCy with Neural Networks.")
    add_para(doc, "[2] Streamlit Documentation. Streamlit: A faster way to build and share data apps.")
    add_para(doc, "[3] PaddleNLP Dataset Documentation. CLUEWSC2020 dataset.")
    add_para(doc, "[4] Jurafsky, D. and Martin, J. H. Speech and Language Processing.")

    doc.save(str(OUT_PATH))
    print(OUT_PATH)


if __name__ == "__main__":
    build_report()
